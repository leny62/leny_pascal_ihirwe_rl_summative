"""Build the submission DOCX and PDF from REPORT.md.

    uv run python -m analysis.build_report

Pandoc converts the markdown with its own default styling, then this script
applies the course template's look on top: the university banner, fonts,
compact spacing, bordered tables and figure sizing.

The template is deliberately not passed to pandoc as --reference-doc. Its table
style does not survive the LibreOffice PDF conversion: cells render empty and
their text spills out underneath as loose paragraphs. Pandoc's own table styling
converts correctly, so the banner and fonts are reapplied here instead.
"""

from __future__ import annotations

import argparse
import copy
import re
import subprocess
import sys
from pathlib import Path

import docx
from docx.enum.section import WD_SECTION_START  # noqa: F401  (kept for clarity)
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parent.parent
MARKDOWN = REPO / "REPORT.md"
TEMPLATE = REPO / "docs/Machine_Learning_Techniques_II - Summative_Assignment - Report Template.docx"
OUT_DOCX = REPO / "assets/report.docx"
OUT_PDF = REPO / "assets/report.pdf"

BODY_FONT = "Times New Roman"  # matches the run font used throughout the template
BODY_PT = 9.5
TABLE_PT = 6.5
HEADING_PT = {1: 15, 2: 12, 3: 10.5}
MAX_FIG = {"w": 5.2, "h": 2.75}


def run_pandoc() -> None:
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        ["pandoc", str(MARKDOWN), "-o", str(OUT_DOCX)],
        cwd=REPO, check=True, capture_output=True, text=True,
    )


def set_page(doc: docx.Document) -> None:
    for s in doc.sections:
        # Pandoc's default output leaves the page size unset, which makes every
        # width calculation below return None.
        s.page_width = Inches(8.27)
        s.page_height = Inches(11.69)
        s.top_margin = Pt(40)
        s.bottom_margin = Pt(36)
        s.left_margin = Pt(54)
        s.right_margin = Pt(54)
        s.header_distance = Pt(20)
        s.footer_distance = Pt(18)


def add_banner(doc: docx.Document) -> bool:
    """Copy the template banner into a first-page-only header."""
    if not TEMPLATE.exists():
        return False
    src = docx.Document(str(TEMPLATE))
    banner = [p for p in src.sections[0].header.paragraphs if p.text.strip()]
    if not banner:
        return False
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    first = sec.first_page_header
    for p in list(first.paragraphs):
        p._p.getparent().remove(p._p)
    for p in banner:
        first._element.append(copy.deepcopy(p._p))
    return True


def add_page_numbers(doc: docx.Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for instr in ("begin", "text", "separate", "end"):
        if instr == "text":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), instr)
        run._r.append(el)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)


def style_text(doc: docx.Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_PT)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(3.5)
    pf.space_before = Pt(0)

    for level, pts in HEADING_PT.items():
        try:
            s = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        s.font.name = BODY_FONT
        s.font.size = Pt(pts)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        s.paragraph_format.space_before = Pt(8 if level > 1 else 0)
        s.paragraph_format.space_after = Pt(3)
        s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        s.paragraph_format.keep_with_next = True

    for p in doc.paragraphs:
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_before = Pt(0)
        is_head = p.style.name.startswith("Heading")
        if not is_head:
            pf.space_after = Pt(3.5)
        # Figure captions are the italic paragraphs; make them small and grey.
        caption = p.runs and all(r.italic for r in p.runs if r.text.strip())
        for r in p.runs:
            r.font.name = BODY_FONT
            if is_head:
                continue
            if caption and p.text.strip().startswith("Figure"):
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            else:
                r.font.size = Pt(BODY_PT)
        if caption and p.text.strip().startswith("Figure"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(7)


def style_tables(doc: docx.Document) -> None:
    for t in doc.tables:
        try:
            t.style = doc.styles["Table Grid"]
        except KeyError:
            pass
        t.autofit = True
        tblPr = t._tbl.tblPr
        old = tblPr.find(qn("w:tblW"))
        if old is not None:
            tblPr.remove(old)
        w = OxmlElement("w:tblW")
        w.set(qn("w:w"), "5000")
        w.set(qn("w:type"), "pct")
        tblPr.append(w)
        marg = tblPr.find(qn("w:tblCellMar")) or OxmlElement("w:tblCellMar")
        if marg.getparent() is None:
            tblPr.append(marg)
        for side, val in (("top", "8"), ("bottom", "8"), ("left", "45"), ("right", "45")):
            el = marg.find(qn(f"w:{side}"))
            if el is None:
                el = OxmlElement(f"w:{side}")
                marg.append(el)
            el.set(qn("w:w"), val)
            el.set(qn("w:type"), "dxa")
        # repeat the header row when a table splits across pages
        trPr = t.rows[0]._tr.get_or_add_trPr()
        if trPr.find(qn("w:tblHeader")) is None:
            trPr.append(OxmlElement("w:tblHeader"))
        for i, row in enumerate(t.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    for r in p.runs:
                        r.font.size = Pt(TABLE_PT)
                        r.font.name = BODY_FONT
                        if i == 0:
                            r.font.bold = True


def strip_bookmarks(doc: docx.Document) -> int:
    """Remove pandoc's per-heading bookmark anchors.

    Pandoc emits a w:bookmarkStart/End pair around every heading so internal
    links resolve. Google Docs draws each one as a blue bookmark ribbon in the
    margin, which looks like a defect in a submitted report. Nothing in this
    document links internally, so the anchors are pure noise.
    """
    removed = 0
    body = doc.element.body
    for tag in ("w:bookmarkStart", "w:bookmarkEnd"):
        for el in body.findall(".//" + qn(tag)):
            el.getparent().remove(el)
            removed += 1
    return removed


def set_default_font(doc: docx.Document) -> None:
    """Force the theme and docDefaults fonts too.

    Setting the Normal style alone is not enough: runs that carry no explicit
    rFonts fall through to docDefaults, and Word/Docs then substitute whatever
    the theme names, which is how a stray second typeface appears.
    """
    dd = doc.styles.element.find(qn("w:docDefaults"))
    if dd is None:
        return
    for rpr in dd.iter(qn("w:rPr")):
        fonts = rpr.find(qn("w:rFonts"))
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            rpr.insert(0, fonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            fonts.set(qn(attr), BODY_FONT)


def drop_alt_text_captions(doc: docx.Document) -> int:
    """Pandoc turns image alt text into its own caption paragraph, so each figure
    gets a bare "Figure 1" line directly above the real caption. Remove those."""
    bare = re.compile(r"^Figure\s+[A-Z0-9]+$")
    removed = 0
    for p in list(doc.paragraphs):
        if bare.match(p.text.strip()) and not p._p.findall(".//" + qn("w:drawing")):
            p._p.getparent().remove(p._p)
            removed += 1
    return removed


def size_figures(doc: docx.Document, max_w: float, max_h: float) -> None:
    mw, mh = Inches(max_w), Inches(max_h)
    for sh in doc.inline_shapes:
        scale = min(mw / sh.width, mh / sh.height, 1.0)
        sh.width = int(sh.width * scale)
        sh.height = int(sh.height * scale)
    # centre the paragraph that holds each image
    for p in doc.paragraphs:
        if p._p.findall(".//" + qn("w:drawing")):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(4)


def to_pdf() -> int | None:
    if OUT_PDF.exists():
        OUT_PDF.unlink()
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir",
         str(OUT_DOCX.parent), str(OUT_DOCX)],
        check=False, capture_output=True,
    )
    if not OUT_PDF.exists():
        return None
    blob = OUT_PDF.read_bytes()
    m = re.search(rb"/Type\s*/Pages[^>]*?/Count\s+(\d+)", blob, re.S)
    return int(m.group(1)) if m else None


def main() -> None:
    ap = argparse.ArgumentParser(description="Build the report DOCX and PDF")
    ap.add_argument("--fig-width", type=float, default=MAX_FIG["w"])
    ap.add_argument("--fig-height", type=float, default=MAX_FIG["h"])
    args = ap.parse_args()

    run_pandoc()
    doc = docx.Document(str(OUT_DOCX))
    set_page(doc)
    banner = add_banner(doc)
    add_page_numbers(doc)
    set_default_font(doc)
    style_text(doc)
    style_tables(doc)
    dropped = drop_alt_text_captions(doc)
    marks = strip_bookmarks(doc)
    size_figures(doc, args.fig_width, args.fig_height)
    doc.save(str(OUT_DOCX))

    pages = to_pdf()
    print(f"wrote {OUT_DOCX.relative_to(REPO)} (banner={'yes' if banner else 'no'}, "
          f"tables={len(doc.tables)}, figures={len(doc.inline_shapes)}, "
          f"alt-captions removed={dropped}, bookmarks stripped={marks})")
    if pages is None:
        print("PDF conversion failed; open the DOCX and export manually")
        sys.exit(1)
    verdict = "OK" if 7 <= pages <= 10 else ("OVER LIMIT" if pages > 10 else "UNDER MINIMUM")
    print(f"wrote {OUT_PDF.relative_to(REPO)}: {pages} pages [{verdict}]")


if __name__ == "__main__":
    main()
